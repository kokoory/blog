# -*- coding:utf-8 -*-
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import time
from time import gmtime, strftime
import datetime

#########
dt = datetime.datetime.today()

#181112
year = dt.strftime("%y")
month =  dt.strftime("%m")
day = dt.strftime("%d")

#이전 달 구하기
first = dt.replace(day=1)
last = first - datetime.timedelta(days=1)

#이전달
prevmonth = last.strftime("%m")

#공문제목
title = "제  목 \'우리제품("+ prevmonth + "월분)\'의 서비스 이용료 지급 요청"

#서비스 청구기간
term = year + ". " + prevmonth + ". 10." + " ~ " + year + ". " + month + ". 9."

#기안일
reportdate = year +". " + month + ". " + day + "."

#기안자
reporter = "기안자 코코리(" + reportdate + ")	   "

#문서번호
docnum = "문서번호 BL-18-D18" + month + day + "01	    "

#문서발신일자
sendingdate = "발신일자 " + reportdate
#########

#공문작성
doc = docx.Document('form.docx')

#변경사항1 제목 > 월
para = doc.paragraphs
para[1].add_run(unicode(title,'utf-8')).bold = True

#아래
para[5].alignment = WD_ALIGN_PARAGRAPH.CENTER

#변경사항2 표 > 청구 기간
table = doc.tables[0]

for row in table.rows:
	for cell in row.cells:
		for para in cell.paragraphs:
			if(para.text == "term"):

				#출력내용 
				para.text = unicode(term,'utf-8')


#변경사항3 기안날짜 
para = doc.paragraphs
para[18].text = unicode(reporter + docnum + sendingdate, 'utf-8')

doc.save('fin_version.docx') 